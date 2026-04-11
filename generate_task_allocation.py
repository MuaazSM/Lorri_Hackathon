"""Generate task allocation Word doc for the optimizer module."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Lorri AI — Optimizer Module Task Allocation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Directory: backend/app/optimizer/')
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph(
    "This document records the division of work across the Operations Research "
    "layer of the Lorri AI freight consolidation platform. Each team member is "
    "responsible for the files, algorithms, and deliverables listed under their "
    "section."
)

# ---------------- Member sections ----------------
members = [
    {
        "name": "Manikya",
        "role": "Heuristic Solver",
        "files": ["heuristic.py"],
        "tasks": [
            "Implement and maintain the First-Fit Decreasing (FFD) bin-packing heuristic used as the fallback solver when shipment count exceeds 50.",
            "Sort shipments by weight (heaviest first) and vehicles by capacity (largest first) to maximize packing efficiency.",
            "Greedy assignment: for each shipment, locate the first vehicle with sufficient weight and volume capacity and full compatibility with all shipments already loaded on that truck.",
            "Apply local-search swap moves between trucks to improve utilization and reduce truck count post-initial-assignment.",
            "Honor the ML-generated compatibility graph — only place shipments together when every pair on a truck is compatible.",
            "Return the standard solver response dict (assigned, unassigned, is_infeasible, plan_metrics) so the LangGraph solver_node can consume it interchangeably with the MIP output.",
            "Ensure O(n²) runtime so large instances (>50 shipments) complete well within pipeline latency budgets.",
        ],
    },
    {
        "name": "Vaishnavi",
        "role": "Queuing Model — Warehouse Congestion",
        "files": ["warehouse_queue.py"],
        "tasks": [
            "Design and implement the queuing-theory model that analyzes warehouse congestion from the final consolidation plan.",
            "Group trucks by origin warehouse and compute arrival/service rates (λ, μ) for each depot based on pickup timestamps and dock throughput.",
            "Compute standard queuing metrics — utilization (ρ), average wait time, queue length, and risk-of-overflow — per warehouse.",
            "Flag bottleneck warehouses where congestion exceeds operational thresholds and expose results to the insight agent for narrative reporting.",
            "Integrate the output into the pipeline so scenario recommendations can trade off consolidation gains against warehouse congestion.",
        ],
    },
    {
        "name": "Muaaz",
        "role": "CP-SAT MIP Solver, Dataset Decision & Task Allocation",
        "files": ["solver.py"],
        "tasks": [
            "Own the core MIP formulation in solver.py using Google OR-Tools CP-SAT.",
            "Define decision variables x[i,k] (shipment → vehicle assignment) and y[k] (vehicle-used indicator).",
            "Encode constraints: exactly-one assignment, per-truck weight capacity, per-truck volume capacity (integer-scaled), y[k] linkage, and pairwise conflict exclusion for forbidden handling combinations and zero time-window overlap.",
            "Formulate the multi-objective function: minimize Σ operating_cost · y[k] − α · Σ utilization, with α tunable for utilization bonus weighting.",
            "Configure CP-SAT solver parameters — 30s time limit, 4 parallel workers — and parse OPTIMAL / FEASIBLE / INFEASIBLE / TIMEOUT outcomes.",
            "Build the solution parser that converts raw CP-SAT values into assignment dicts with per-truck utilization and route detour estimates.",
            "Dataset decision: choose and justify the datasets powering the platform — synthetic Indian freight data under data/synthetic/ (runtime) and the Solomon VRPTW benchmark under dataset/ (academic validation, C1/C2/R1/R2/RC1/RC2 instance families).",
            "Task allocation: divide the optimizer workload among the team and track deliverables (this document).",
        ],
    },
    {
        "name": "Aditya",
        "role": "Load Consolidation — Compatibility, Routing, Metrics & Sensitivity",
        "files": [
            "compatibility.py",
            "route_optimizer.py",
            "metrics.py",
            "baseline.py",
            "sensitivity.py",
            "__init__.py",
        ],
        "tasks": [
            "Own the end-to-end load consolidation layer that wraps the solver output into business-ready plans.",
            "compatibility.py — apply hard rule-based filters on the ML compatibility graph: minimum time-window overlap, combined weight+volume vs max vehicle capacity, forbidden handling pairs (hazmat+fragile, etc.), and maximum route detour (500 km default for Indian lanes).",
            "route_optimizer.py — after assignment, run a TSP per truck using OR-Tools RoutingModel (PATH_CHEAPEST_ARC + GUIDED_LOCAL_SEARCH) to compute the optimal visit sequence of origin/destination cities; fall back to brute-force permutation for ≤8 cities.",
            "baseline.py — compute the one-truck-per-shipment baseline used as the before/after reference for savings reporting.",
            "metrics.py — compute full before/after KPIs: trip reduction %, cost savings %, carbon savings %, fleet utilization, and per-truck detail.",
            "sensitivity.py — run robustness analysis by perturbing cost/capacity/demand inputs and reporting plan stability.",
            "__init__.py — maintain the clean public API (solve, compute_metrics, compute_baseline_metrics, optimize_all_routes, run_sensitivity_analysis, analyze_warehouse_congestion) that the LangGraph nodes, API routes, and tests import from.",
        ],
    },
]

for m in members:
    doc.add_heading(f"{m['name']} — {m['role']}", level=1)

    p = doc.add_paragraph()
    p.add_run("Files owned: ").bold = True
    p.add_run(", ".join(m["files"]))

    p = doc.add_paragraph()
    p.add_run("Tasks completed:").bold = True
    for task in m["tasks"]:
        doc.add_paragraph(task, style="List Bullet")

# Summary table
doc.add_heading("Ownership Summary", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Member"
hdr[1].text = "Area"
hdr[2].text = "Files"
for m in members:
    row = table.add_row().cells
    row[0].text = m["name"]
    row[1].text = m["role"]
    row[2].text = ", ".join(m["files"])

out_path = "/Users/muaazshaikh/Lorri_Hackathon/Optimizer_Task_Allocation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
